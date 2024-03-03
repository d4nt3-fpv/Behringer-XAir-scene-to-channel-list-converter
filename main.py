from tkinter import *
from tkinter.ttk import *
from tkinter import messagebox
from tkinter import filedialog
import docx
from docx2pdf import convert
  

bus_names_global_var = []

def readSceneFile(filename):
    start_with_list = ["/ch/01/config", "/ch/02/config", "/ch/03/config", "/ch/04/config", "/ch/05/config", "/ch/06/config", "/ch/07/config", "/ch/08/config", "/ch/09/config", "/ch/10/config", "/ch/11/config", "/ch/12/config", "/ch/13/config", "/ch/14/config", "/ch/15/config", "/ch/16/config", "/rtn/aux/config"]

    start_with_list_busses = ["/bus/1/config", "/bus/2/config", "/bus/3/config", "/bus/4/config", "/bus/5/config", "/bus/6/config"]

    file_to_read = filename

    interessting_lines = []
    interessting_lines_busses = []

    channel_names = []
    bus_names = []

    f = open(file_to_read, "r")

    for line in f:
        # print(line)

        for i in start_with_list:
            if line.startswith(i):
                interessting_lines.append(line)

        for z in start_with_list_busses:
            if line.startswith(z):
                interessting_lines_busses.append(line)

    # print(interessting_lines)

    for element in interessting_lines:
        new_element_sliced = element.split('"')
        # print(new_element_sliced[1])
        channel_names.append(new_element_sliced[1])

    for item in interessting_lines_busses:
        new_item_sliced = item.split('"')
        bus_names.append(new_item_sliced[1])

    returning_data = []
    global bus_names_global_var
    bus_names_global_var = []
    i = 0

    for channel_name in channel_names:
        i += 1
        returning_data.append((i, channel_name))
    returning_data.append((18, channel_names[16]))
    print(returning_data)
    
    y = 0
    for bus_name in bus_names:
        y += 1
        bus_names_global_var.append(("Bus " + str(y), bus_name))

    print(bus_names_global_var)

    return returning_data





def CreateWordDocument(filename, title ,tabledata, headding_row1, headding_row2, include_busses, bus_table_data):
    # Create an instance of a word document 
    doc = docx.Document()  

    # Add a Title to the document 
    doc.add_heading(title, 0) 
  
    # Creating a table object 
    table = doc.add_table(rows=1, cols=2) 
  
    # Adding heading in the 1st row of the table 
    row = table.rows[0].cells 
    row[0].text = headding_row1
    row[1].text = headding_row2

    data = tabledata
  
    # Adding data from the list to the table 
    for id, name in data: 
    
        # Adding a row and then adding data in it. 
        row = table.add_row().cells 
        # Converting id to string as table can only take string input 
        row[0].text = str(id) 
        row[1].text = name 

    table.style = 'Colorful List'


    if include_busses:
        doc.add_heading('Bus sends:', 1)



        # Creating a table object 
        bus_table = doc.add_table(rows=1, cols=2) 
    
        # Adding heading in the 1st row of the table 
        row = bus_table.rows[0].cells 
        row[0].text = "Nr."
        row[1].text = "Name"

        bus_data = bus_table_data
    
        # Adding data from the list to the table 
        for id, name in bus_data: 
        
            # Adding a row and then adding data in it. 
            row = bus_table.add_row().cells 
            # Converting id to string as table can only take string input 
            row[0].text = id
            row[1].text = name 

        bus_table.style = 'Colorful List'

    # Now save the document to a location 
    doc.save(filename)
    messagebox.showinfo(title="Success", message="Success.")


# converttopdf = False

# CreateWordDocument("channellist.docx", "Recknitzgala Inputs", readSceneFile("test.scn"), "Channel", "Name")

# if converttopdf:
#     convert("channellist.docx")

current_channel_data = []



def loadfilebtnclick():
        
    folder = filedialog.askopenfilename(defaultextension=".scn", filetypes=[("Scene", "*.scn")])
    print(folder)
    global current_channel_data
    current_channel_data = []
    current_channel_data = readSceneFile(folder)
    global Button_0
    Button_0.config(state=NORMAL)


def savefilebtnclick():
    folder = filedialog.asksaveasfilename(defaultextension=".docx")
    print(folder)
    global current_channel_data
    print(current_channel_data)
    if (checkbox_var.get() == 1):
        include_busses_bool= True
    else:
        include_busses_bool = False
    CreateWordDocument(folder, "Channels", current_channel_data, "Nr.", "Name", include_busses_bool, bus_names_global_var)


window=Tk()
window.title('Scene to Channel list for Behringer XR18')
window.geometry('251x300+10+20')
window.resizable(False, False)

checkbox_var = IntVar()
checkbox_var.set(0)


########### Labels ###########
# Label_0 = Label(window, text='Video URL').place(x=32,y=98)
Label_1 = Label(window, text='Convert XR18 Scene to channel list').place(x=15,y=44)



########### Checkboxes ###########
Checkbox_0 = Checkbutton(window, text='Include bus sends', variable=checkbox_var).place(x=32,y=145)


########### Buttons ###########
Button_1 = Button(window, text='Load scene file...', command=loadfilebtnclick)
Button_1.place(x=60,y=98)
Button_0 = Button(window, text='Save docx file...',command=savefilebtnclick, state=DISABLED)
Button_0.place(x=60,y=192)

########### Entrys ###########
# Entry_0 = Entry(window, text='Entry_0').place(x=95,y=97)



window.mainloop()

